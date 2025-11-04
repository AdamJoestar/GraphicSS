import sys
from PyQt5.QtWidgets import (
    QApplication,
    QWidget,
    QLabel,
    QVBoxLayout,
    QPushButton,
    QMessageBox,
    QDialog,
    QLineEdit,
    QInputDialog,
    QFileDialog, # New: For file selection
)
from PyQt5.QtGui import QPixmap, QPainter, QPen, QColor, QImage
from PyQt5.QtCore import Qt, QPoint, QRect, QDir # New: QDir
from PIL import ImageGrab, Image
from docx import Document
from docx.shared import Inches
from datetime import datetime
import os
import time

# --- Configuration & Template Management ---

def create_dummy_template(filepath, placeholder_text):
    """Creates a basic template with the required placeholder if the file does not exist."""
    print(f"Template '{filepath}' not found. Creating a dummy template with placeholder.")
    doc = Document()
    doc.add_heading("INTERACTIVE TEST REPORT", 0)
    doc.add_paragraph("This report was generated using the screenshot tool.")
    doc.add_heading("Results Insertion Area:", level=3)
    doc.add_paragraph(f"Placeholder for content: **{placeholder_text}**")
    doc.add_paragraph(placeholder_text) # The actual placeholder string
    doc.save(filepath)
    print(f"Dummy Word template '{filepath}' created.")
    return filepath


class ScreenshotSelector(QDialog):
    """A frameless dialog to allow the user to select an area of the screen for cropping."""

    def __init__(self, full_screen_image, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Select Screenshot Area")
        # Set the dialog size to cover the entire screen
        desktop = QApplication.instance().desktop()
        self.setGeometry(0, 0, desktop.width(), desktop.height())
        # Frameless window, always on top
        self.setWindowFlags(Qt.FramelessWindowHint | Qt.WindowStaysOnTopHint)
        self.setCursor(Qt.CrossCursor)  # Cursor becomes crosshair

        self.full_screen_pixmap = QPixmap.fromImage(full_screen_image)
        self.start_point = QPoint()
        self.end_point = QPoint()
        self.selection_rect = QRect()
        self.is_drawing = False

    def paintEvent(self, event):
        painter = QPainter(self)
        # Draw the full screen background image
        painter.drawPixmap(0, 0, self.full_screen_pixmap)

        # Draw the selection rectangle if currently drawing
        if self.is_drawing:
            pen = QPen(QColor(255, 0, 0))  # Red color
            pen.setWidth(2)
            painter.setPen(pen)
            painter.drawRect(self.selection_rect)

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.start_point = event.pos()
            self.end_point = event.pos()
            self.is_drawing = True
            self.selection_rect = QRect(self.start_point, self.end_point).normalized()
            self.update()

    def mouseMoveEvent(self, event):
        if self.is_drawing:
            self.end_point = event.pos()
            self.selection_rect = QRect(self.start_point, self.end_point).normalized()
            self.update()

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.is_drawing = False
            self.close()  # Close window after selection is complete

class App:
    """Main application logic for capturing screenshots and generating the report."""

    def __init__(self):
        self.screenshots = []  # List to store multiple screenshots and their info
        self.input_template_path = ""
        self.final_report_name = ""
        self.final_pdf_name = ""
        self.base_placeholder = "[GRAPH_{0}]"  # Placeholder template with index

    def get_report_details(self):
        """Prompts the user for all necessary report details."""
        
        # 1. Get Number of Screenshots
        num_screenshots, ok = QInputDialog.getInt(
            None,
            "Number of Screenshots",
            "How many screenshots do you want to insert?",
            1, 1, 10  # min 1, max 10 screenshots
        )
        if not ok:
            return False
            
        # Store the placeholders
        self.screenshots = []
        for i in range(num_screenshots):
            placeholder = self.base_placeholder.format(i + 1)
            self.screenshots.append({
                'placeholder': placeholder,
                'image_path': f"graph_temp_{i+1}.png",
                'custom_text': f"Figure {i+1}"
            })

        # 2. Get Input Template Path (using QFileDialog for better UX)
        input_template_path, _ = QFileDialog.getOpenFileName(
            None,
            "Select Input Word Template (.docx)",
            QDir.currentPath(),
            "Word Documents (*.docx)",
        )
        if input_template_path:
            self.input_template_path = input_template_path
        else:
            QMessageBox.warning(None, "Warning", "No template file selected. Using 'default_template.docx'.")
            self.input_template_path = "default_template.docx"

        # If the selected file doesn't exist, create a dummy one
        if not os.path.exists(self.input_template_path):
            # use the first placeholder as the template placeholder
            create_dummy_template(self.input_template_path, self.base_placeholder.format(1))


        # 3. Get Final Output Report Name
        base_template_name = os.path.splitext(os.path.basename(self.input_template_path))[0]
        default_output_name = f"{base_template_name}_Output_{datetime.now().strftime('%Y%m%d')}.docx"

        # Use File Save Dialog to get both the path and name
        self.final_report_name, _ = QFileDialog.getSaveFileName(
            None,
            "Save Report As",
            os.path.join(QDir.homePath(), default_output_name),  # Start in user's home directory
            "Word Documents (*.docx)"
        )

        if self.final_report_name:
            if not self.final_report_name.lower().endswith(".docx"):
                self.final_report_name += ".docx"
            
            # Set PDF name in the same directory as the DOCX
            base_name = os.path.splitext(self.final_report_name)[0]
            self.final_pdf_name = base_name + ".pdf"
            return True
        else:
            return False

    def take_interactive_screenshot(self, prompt_text="Select area"):
        """Initiates full-screen capture and interactive selection."""
        QMessageBox.information(
            None,
            "Ready for Screenshot",
            f"Click OK, then **{prompt_text}** using Drag-and-Drop on the screen.",
        )

        app = QApplication.instance()
        if app:
            for widget in app.topLevelWidgets():
                # Check if it's the main application window and hide it
                if isinstance(widget, QWidget) and widget.windowTitle() == "PyQt Report Generator":
                    widget.hide()
                    break # Assuming only one main window to hide

        time.sleep(0.5) 

        full_screen_image_pil = ImageGrab.grab()

        if app:
            for widget in app.topLevelWidgets():
                if isinstance(widget, QWidget) and widget.windowTitle() == "PyQt Report Generator":
                    widget.show()
                    break

        # Convert PIL Image to Qt-compatible format (QImage)
        img_data = full_screen_image_pil.tobytes("raw", "RGB")
        full_screen_image_qt = QImage(
            img_data,
            full_screen_image_pil.size[0],
            full_screen_image_pil.size[1],
            full_screen_image_pil.size[0] * 3,
            QImage.Format_RGB888,
        )

        selector = ScreenshotSelector(full_screen_image_qt)
        selector.exec()

        x = selector.selection_rect.x()
        y = selector.selection_rect.y()
        w = selector.selection_rect.width()
        h = selector.selection_rect.height()

        if w == 0 or h == 0:
            QMessageBox.warning(None, "Warning", "Area not selected or too small. Please try again.")
            return None

        # Crop image according to selection
        cropped_image_pil = full_screen_image_pil.crop((x, y, x + w, y + h))
        return cropped_image_pil

    def insert_content_at_placeholder(self, document, date_str):
        """Insert all screenshots and their custom text into the document at their placeholders.

        Returns (document, success)
        """
        for screenshot in self.screenshots:
            insertion_found = False
            for paragraph in document.paragraphs:
                if screenshot['placeholder'] in paragraph.text:
                    # Clear the placeholder
                    paragraph.text = paragraph.text.replace(screenshot['placeholder'], "")

                    # Ask for custom text (default to stored custom_text)
                    custom_text, ok = QInputDialog.getText(
                        None,
                        "Custom Text Input",
                        f"Enter the text you want to add before the date for {screenshot['placeholder']}:",
                        QLineEdit.Normal,
                        screenshot.get('custom_text', 'Report Content Inserted')
                    )
                    if not ok or not custom_text:
                        custom_text = screenshot.get('custom_text', 'Report Content Inserted')

                    # Insert custom text + date
                    run_date = paragraph.add_run()
                    run_date.add_text(f"{custom_text}: {date_str}\n")
                    run_date.bold = True

                    # Insert image
                    try:
                        paragraph.add_run().add_picture(screenshot['image_path'], width=Inches(6))
                    except Exception as e:
                        QMessageBox.critical(None, "Error", f"Failed to insert image {screenshot['image_path']}: {e}")
                        return document, False

                    insertion_found = True
                    break

            if not insertion_found:
                QMessageBox.critical(None, "Error", f"Could not find the unique placeholder text: '{screenshot['placeholder']}' in the document.")
                return document, False

        return document, True

    def run(self):
        # 0. Set up a dummy main window title for hiding logic
        main_window = QWidget()
        main_window.setWindowTitle("PyQt Report Generator")
        main_window.setGeometry(100, 100, 300, 100)
        main_window.show()


        # 1. Get Report Details from User
        if not self.get_report_details():
            QMessageBox.critical(None, "Error", "Report details not provided. Exiting.")
            main_window.close()
            return

        # show all placeholders that will be used
        placeholders_str = ", ".join([s['placeholder'] for s in self.screenshots])
        QMessageBox.information(
            None,
            "Start Application",
            f"Input Template: **{os.path.basename(self.input_template_path)}**\nOutput Report: **{self.final_report_name}**\nInsertion Placeholders: **{placeholders_str}**\n\nClick OK to start screenshot capture.",
        )
        main_window.hide()

        # 2. Capture Multiple Screenshots (with preview and retake)
        for i, screenshot_info in enumerate(self.screenshots):
            QMessageBox.information(
                None,
                "Next Screenshot",
                f"Ready to capture screenshot {i+1} of {len(self.screenshots)}\nThis will be inserted at placeholder: {screenshot_info['placeholder']}"
            )

            while True:
                graph_image = self.take_interactive_screenshot(f"select area for {screenshot_info['placeholder']}")

                if not graph_image:
                    main_window.show()
                    return

                # Save temporary image for preview
                graph_image.save(screenshot_info['image_path'])
                print(f"Screenshot {i+1} captured.")

                # Create preview dialog
                preview_dialog = QDialog()
                preview_dialog.setWindowTitle("Screenshot Preview")
                preview_dialog.setMinimumSize(400, 400)

                # Create layout
                layout = QVBoxLayout()

                # Add preview image
                preview_label = QLabel()
                pixmap = QPixmap(screenshot_info['image_path'])
                scaled_pixmap = pixmap.scaled(350, 350, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                preview_label.setPixmap(scaled_pixmap)
                layout.addWidget(preview_label)

                # Add confirmation message
                layout.addWidget(QLabel("Is this screenshot correct?"))

                # Add buttons
                button_box = QVBoxLayout()
                confirm_button = QPushButton("Yes, Continue")
                retake_button = QPushButton("No, Retake Screenshot")
                cancel_button = QPushButton("Cancel")

                button_box.addWidget(confirm_button)
                button_box.addWidget(retake_button)
                button_box.addWidget(cancel_button)
                layout.addLayout(button_box)

                preview_dialog.setLayout(layout)

                # Connect button signals
                confirm_button.clicked.connect(preview_dialog.accept)
                retake_button.clicked.connect(preview_dialog.reject)
                cancel_button.clicked.connect(lambda: preview_dialog.done(2))

                # Show dialog and get result
                result = preview_dialog.exec_()

                if result == QDialog.Accepted:  # User confirmed screenshot
                    print("Screenshot confirmed.")
                    break
                elif result == 2:  # User cancelled
                    # remove any saved temp files and exit
                    if os.path.exists(screenshot_info['image_path']):
                        os.remove(screenshot_info['image_path'])
                    main_window.show()
                    return
                else:  # User wants to retake
                    print("Retaking screenshot...")
                    continue

        # 3. Create Word Report
        print("Starting Word document creation...")
        try:
            document = Document(self.input_template_path)
        except Exception as e:
            QMessageBox.critical(None, "Error Loading Document", f"Failed to load document '{self.input_template_path}'. Error: {e}")
            main_window.show()
            return

        # Get the current date
        current_date_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Insert all screenshots and their content into the document
        document, success = self.insert_content_at_placeholder(document, current_date_str)
        
        if not success:
            main_window.show()
            return

        document.save(self.final_report_name)
        print(f"\n✅ Final Word report saved as: {self.final_report_name}")

        # 4. Convert to PDF (Optional/Non-Essential)
        try:
            import docxtopdf
            docxtopdf.convert(self.final_report_name, self.final_pdf_name)
            print(f"✅ Final PDF report saved as: {self.final_pdf_name}")
        except ImportError:
            print("Warning: docxtopdf is not installed or does not support this OS. PDF not created.")
        except Exception as e:
            print(f"Error converting to PDF: {e}. Make sure MS Word is installed.")

        # 5. Clean up temporary files
        for screenshot in self.screenshots:
            if os.path.exists(screenshot['image_path']):
                os.remove(screenshot['image_path'])
        print("Temporary files cleaned up.")
        
        main_window.show()
        QMessageBox.information(
            None,
            "Finished",
            f"Report '{self.final_report_name}' has been created from '{os.path.basename(self.input_template_path)}'.",
        )


if __name__ == '__main__':
    # Initialize the QApplication instance before creating any Qt widgets
    app = QApplication(sys.argv)
    main_app = App()
    
    while True:
        main_app.run()
        
        # Ask if user wants to continue
        response = QMessageBox.question(
            None,
            "Continue?",
            "Do you want to perform another task?",
            QMessageBox.Yes | QMessageBox.No
        )
        
        if response == QMessageBox.No:
            break
    
    # Clean up and exit
    app.quit()
    sys.exit()